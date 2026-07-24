from __future__ import annotations

import io
from pathlib import Path
from typing import Any

import docx

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_PJ_TEMPLATE = ROOT / "assets" / "templates" / "proposta_comercial_verdio.docx"


def _replace_in_paragraph(paragraph: Any, replacements: dict[str, str]) -> None:
    full_text = paragraph.text
    if not any(token in full_text for token in replacements):
        return

    updated = full_text
    for token, value in replacements.items():
        updated = updated.replace(token, value)

    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = updated


def generate_pj_proposal(
    context: dict[str, Any],
    *,
    template_path: str | Path | None = None,
) -> bytes:
    """Gera o DOCX comercial sem expor custos ou margens internas."""
    path = Path(template_path) if template_path else DEFAULT_PJ_TEMPLATE
    if not path.exists():
        raise FileNotFoundError(f"Template não encontrado: {path}")

    document = docx.Document(str(path))
    replacements = {
        "{{" + key + "}}": str(value)
        for key, value in context.items()
        if key != "itens_proposta"
    }

    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    if document.tables:
        product_table = document.tables[0]
        for item in context.get("itens_proposta", []):
            cells = product_table.add_row().cells
            if len(cells) >= 3:
                cells[0].text = str(item.get("nome", ""))
                cells[1].text = str(item.get("descricao", ""))
                cells[2].text = str(item.get("preco", ""))

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
