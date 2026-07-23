from __future__ import annotations

import hashlib
import io
from datetime import date, datetime, timedelta
from decimal import Decimal
from pathlib import Path

import docx
import pandas as pd
import streamlit as st

import user_management_db as db
from app_core.auth import require_auth
from app_core.pricing import (
    gross_margin_percent,
    gross_margin_value,
    quantize_money,
    sale_price_from_margin,
    to_decimal,
)
from app_core.ui import apply_branding, configure_page, money, render_hero, render_sidebar

configure_page("Simulador Pessoa Jurídica")
apply_branding()
require_auth()
render_sidebar()
render_hero(
    "Simulador de venda — Pessoa Jurídica",
    "Configure a frota, o prazo contratual, os produtos e as condições comerciais da proposta.",
)

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT / "assets" / "templates" / "proposta_comercial_verdio.docx"

pricing_config = db.get_pricing_config()
plans = {
    plan: {product: quantize_money(value) for product, value in products.items()}
    for plan, products in pricing_config.get("PLANOS_PJ", {}).items()
}
costs_by_plan = {
    plan: {product: quantize_money(value) for product, value in products.items()}
    for plan, products in pricing_config.get("CUSTOS_PJ", {}).items()
}
descriptions = pricing_config.get("PRODUTOS_PJ_DESCRICAO", {})
is_admin = st.session_state.get("role") == "admin"

if "pj_results" not in st.session_state:
    st.session_state.pj_results = None


def _product_key(product: str) -> str:
    return hashlib.sha1(product.encode("utf-8")).hexdigest()[:10]


def _margin_label(percent: Decimal | None) -> str:
    return "Não disponível" if percent is None else f"{percent:.2f}%"


def replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    full_text = paragraph.text
    if not any(token in full_text for token in replacements):
        return
    updated = full_text
    for token, value in replacements.items():
        updated = updated.replace(token, value)
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = updated


@st.cache_data(show_spinner=False)
def generate_proposal(context: dict) -> bytes:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template não encontrado: {TEMPLATE_PATH}")

    document = docx.Document(str(TEMPLATE_PATH))
    replacements = {
        "{{" + key + "}}": str(value)
        for key, value in context.items()
        if key != "itens_proposta"
    }

    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)

    if document.tables:
        product_table = document.tables[0]
        for item in context.get("itens_proposta", []):
            cells = product_table.add_row().cells
            if len(cells) >= 3:
                cells[0].text = item["nome"]
                cells[1].text = item["descricao"]
                cells[2].text = item["preco"]

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def clear_simulation() -> None:
    keys = [
        key
        for key in list(st.session_state)
        if key.startswith("pj_") and key != "pj_results"
    ]
    for key in keys:
        st.session_state.pop(key, None)
    st.session_state.pj_results = None


if not plans:
    st.warning("Não há planos PJ configurados. Solicite ao administrador a inclusão dos preços.")
    st.stop()

header_actions = st.columns([5, 1])
with header_actions[1]:
    if st.button("Limpar simulação", width="stretch"):
        clear_simulation()
        st.rerun()

config_col, client_col = st.columns([1, 1.4])
with config_col:
    st.markdown("#### Configuração comercial")
    vehicle_count = st.number_input(
        "Quantidade de veículos",
        min_value=1,
        value=1,
        step=1,
        key="pj_vehicle_count",
    )
    contract_term = st.selectbox(
        "Prazo do contrato",
        list(plans),
        key="pj_contract_term",
    )
    st.caption("Os preços e custos são mensais e unitários por veículo.")

with client_col:
    st.markdown("#### Dados da proposta")
    company = st.text_input("Empresa", key="pj_company")
    responsible = st.text_input("Responsável", key="pj_responsible")
    validity_days = st.number_input(
        "Validade da proposta (dias)",
        min_value=1,
        max_value=90,
        value=15,
        key="pj_validity_days",
    )

st.markdown("#### Produtos e serviços")
if is_admin:
    st.info(
        "Como administrador, você visualiza custos e margens e pode substituir o preço "
        "de cada produto por um valor final ou por uma margem-alvo. Esses dados internos "
        "não são exibidos no documento da proposta."
    )

selected: dict[str, dict[str, Decimal | str | bool | None]] = {}
current_plan_costs = costs_by_plan.get(contract_term, {})

for product_index, (product, base_price) in enumerate(plans[contract_term].items()):
    product_id = _product_key(product)
    cost = quantize_money(current_plan_costs.get(product, 0))
    cost_configured = cost > 0
    base_margin_value = gross_margin_value(base_price, cost) if cost_configured else None
    base_margin_percent = gross_margin_percent(base_price, cost) if cost_configured else None

    with st.container(border=True):
        enable_col, value_col = st.columns([1.65, 1])
        with enable_col:
            enabled = st.toggle(
                product,
                key=f"pj_enabled_{contract_term}_{product_id}",
            )
            st.caption(descriptions.get(product, product))
        with value_col:
            st.markdown(f"**Preço padrão: {money(base_price)}**")
            st.caption("por veículo/mês")

        effective_price = base_price
        pricing_mode = "Preço padrão"

        if is_admin:
            margin_1, margin_2, margin_3 = st.columns(3)
            margin_1.metric("Custo mensal", money(cost) if cost_configured else "Não cadastrado")
            margin_2.metric(
                "Margem unitária",
                money(base_margin_value) if base_margin_value is not None else "Pendente",
            )
            margin_3.metric("Margem percentual", _margin_label(base_margin_percent))

            if not cost_configured:
                st.warning(
                    "O custo mensal deste produto ainda não foi cadastrado em "
                    "Administração → Preços e produtos. A margem percentual não pode ser calculada."
                )

            if enabled:
                modes = ["Preço padrão", "Valor personalizado"]
                if cost_configured:
                    modes.append("Margem personalizada")

                pricing_mode = st.selectbox(
                    "Condição comercial",
                    modes,
                    key=f"pj_mode_{contract_term}_{product_id}",
                    help=(
                        "Valor personalizado define o preço mensal final. Margem personalizada "
                        "calcula o preço necessário com base no custo cadastrado."
                    ),
                )

                if pricing_mode == "Valor personalizado":
                    custom_value = st.number_input(
                        "Preço mensal personalizado por veículo",
                        min_value=0.0,
                        value=float(base_price),
                        step=1.0,
                        format="%.2f",
                        key=f"pj_custom_value_{contract_term}_{product_id}",
                    )
                    effective_price = quantize_money(custom_value)
                elif pricing_mode == "Margem personalizada":
                    default_margin = float(base_margin_percent or Decimal("30"))
                    default_margin = max(0.0, min(default_margin, 99.0))
                    target_margin = st.number_input(
                        "Margem desejada sobre o preço de venda (%)",
                        min_value=0.0,
                        max_value=99.0,
                        value=default_margin,
                        step=0.5,
                        format="%.2f",
                        key=f"pj_custom_margin_{contract_term}_{product_id}",
                    )
                    effective_price = sale_price_from_margin(cost, target_margin)
                    st.caption(f"Preço calculado: {money(effective_price)} por veículo/mês")

                if cost_configured:
                    effective_margin_value = gross_margin_value(effective_price, cost)
                    effective_margin_percent = gross_margin_percent(effective_price, cost)
                    if effective_margin_value < 0:
                        st.error(
                            f"O preço selecionado gera prejuízo de {money(abs(effective_margin_value))} "
                            "por veículo/mês."
                        )
                    else:
                        st.success(
                            f"Margem simulada: {money(effective_margin_value)} "
                            f"({_margin_label(effective_margin_percent)}) por veículo/mês."
                        )

        if enabled:
            margin_value = gross_margin_value(effective_price, cost) if cost_configured else None
            margin_percent = gross_margin_percent(effective_price, cost) if cost_configured else None
            selected[product] = {
                "price": effective_price,
                "cost": cost,
                "cost_configured": cost_configured,
                "margin_value": margin_value,
                "margin_percent": margin_percent,
                "pricing_mode": pricing_mode,
            }

submitted = st.button(
    "Calcular e registrar proposta",
    type="primary",
    width="stretch",
    key="pj_submit",
)

if submitted:
    if not company.strip() or not responsible.strip():
        st.warning("Informe a empresa e o responsável.")
    elif not selected:
        st.warning("Selecione pelo menos um produto ou serviço.")
    else:
        vehicle_decimal = Decimal(vehicle_count)
        months = Decimal(contract_term.split()[0])
        monthly_per_vehicle = sum(
            (to_decimal(item["price"]) for item in selected.values()),
            Decimal("0"),
        )
        monthly_fleet = quantize_money(monthly_per_vehicle * vehicle_decimal)
        contract_total = quantize_money(monthly_fleet * months)

        all_costs_configured = all(bool(item["cost_configured"]) for item in selected.values())
        monthly_cost_per_vehicle = sum(
            (to_decimal(item["cost"]) for item in selected.values()),
            Decimal("0"),
        )
        monthly_cost_fleet = quantize_money(monthly_cost_per_vehicle * vehicle_decimal)
        monthly_margin_fleet = quantize_money(monthly_fleet - monthly_cost_fleet)
        contract_margin = quantize_money(monthly_margin_fleet * months)
        aggregate_margin_percent = (
            gross_margin_percent(monthly_per_vehicle, monthly_cost_per_vehicle)
            if all_costs_configured
            else None
        )

        validity_label = (date.today() + timedelta(days=int(validity_days))).strftime("%d/%m/%Y")
        item_rows = []
        proposal_items = []
        for product, item in selected.items():
            price = to_decimal(item["price"])
            cost = to_decimal(item["cost"])
            margin_value = item["margin_value"]
            margin_percent = item["margin_percent"]
            proposal_items.append(
                {
                    "nome": product,
                    "descricao": descriptions.get(product, product),
                    "preco": money(price),
                }
            )
            item_rows.append(
                {
                    "Produto": product,
                    "Condição": str(item["pricing_mode"]),
                    "Preço mensal unitário": float(price),
                    "Custo mensal unitário": float(cost) if item["cost_configured"] else None,
                    "Margem unitária": float(margin_value) if margin_value is not None else None,
                    "Margem (%)": float(margin_percent) if margin_percent is not None else None,
                    "Margem mensal da frota": (
                        float(quantize_money(to_decimal(margin_value) * vehicle_decimal))
                        if margin_value is not None
                        else None
                    ),
                }
            )

        context = {
            "NOME_EMPRESA": company.strip(),
            "NOME_RESPONSAVEL": responsible.strip(),
            "NOME_CONSULTOR": st.session_state.get("name", ""),
            "DATA_VALIDADE": validity_label,
            "QTD_VEICULOS": str(vehicle_count),
            "TEMPO_CONTRATO": contract_term,
            "VALOR_MENSAL_FROTA": money(monthly_fleet),
            "VALOR_TOTAL_CONTRATO": money(contract_total),
            "SOMA_TOTAL_MENSAL_VEICULO": money(monthly_per_vehicle),
            "itens_proposta": proposal_items,
        }
        st.session_state.pj_results = {
            "monthly_per_vehicle": quantize_money(monthly_per_vehicle),
            "monthly_fleet": monthly_fleet,
            "contract_total": contract_total,
            "monthly_cost_per_vehicle": quantize_money(monthly_cost_per_vehicle),
            "monthly_cost_fleet": monthly_cost_fleet,
            "monthly_margin_fleet": monthly_margin_fleet,
            "contract_margin": contract_margin,
            "aggregate_margin_percent": aggregate_margin_percent,
            "all_costs_configured": all_costs_configured,
            "item_rows": item_rows,
            "context": context,
        }
        db.upsert_proposal(
            {
                "tipo": "PJ",
                "empresa": company.strip(),
                "consultor": st.session_state.get("name", "N/A"),
                "valor_total": float(contract_total),
            }
        )
        db.add_log(
            st.session_state.get("username", "sistema"),
            "Simulou proposta PJ",
            {
                "empresa": company.strip(),
                "veiculos": vehicle_count,
                "prazo": contract_term,
                "valor_total": float(contract_total),
                "produtos": [
                    {
                        "produto": product,
                        "condicao": str(item["pricing_mode"]),
                        "preco_mensal": float(to_decimal(item["price"])),
                    }
                    for product, item in selected.items()
                ],
            },
        )
        st.success("Proposta calculada e registrada.")

result = st.session_state.get("pj_results")
if result:
    st.markdown("### Resultado")
    metric_1, metric_2, metric_3 = st.columns(3)
    metric_1.metric("Mensal por veículo", money(result["monthly_per_vehicle"]))
    metric_2.metric("Mensal da frota", money(result["monthly_fleet"]))
    metric_3.metric("Valor do contrato", money(result["contract_total"]))

    if is_admin:
        st.markdown("#### Rentabilidade interna")
        if result["all_costs_configured"]:
            margin_1, margin_2, margin_3, margin_4 = st.columns(4)
            margin_1.metric("Custo mensal da frota", money(result["monthly_cost_fleet"]))
            margin_2.metric("Margem mensal da frota", money(result["monthly_margin_fleet"]))
            margin_3.metric("Margem do contrato", money(result["contract_margin"]))
            margin_4.metric(
                "Margem percentual",
                _margin_label(result["aggregate_margin_percent"]),
            )
        else:
            st.warning(
                "A rentabilidade consolidada não é definitiva porque um ou mais produtos "
                "selecionados estão sem custo cadastrado."
            )

        margin_df = pd.DataFrame(result["item_rows"])
        st.dataframe(
            margin_df,
            width="stretch",
            hide_index=True,
            column_config={
                "Produto": "Produto",
                "Condição": "Condição comercial",
                "Preço mensal unitário": st.column_config.NumberColumn(
                    "Preço mensal unitário",
                    format="R$ %.2f",
                ),
                "Custo mensal unitário": st.column_config.NumberColumn(
                    "Custo mensal unitário",
                    format="R$ %.2f",
                ),
                "Margem unitária": st.column_config.NumberColumn(
                    "Margem unitária",
                    format="R$ %.2f",
                ),
                "Margem (%)": st.column_config.NumberColumn(
                    "Margem (%)",
                    format="%.2f%%",
                ),
                "Margem mensal da frota": st.column_config.NumberColumn(
                    "Margem mensal da frota",
                    format="R$ %.2f",
                ),
            },
        )

    try:
        document_bytes = generate_proposal(result["context"])
        safe_company = "_".join(result["context"]["NOME_EMPRESA"].split())
        st.download_button(
            "Baixar proposta em DOCX",
            data=document_bytes,
            file_name=f"Proposta_{safe_company}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
        )
    except Exception as exc:
        st.error(f"Não foi possível gerar o documento: {exc}")
