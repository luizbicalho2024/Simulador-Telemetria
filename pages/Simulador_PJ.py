# pages/Simulador_PJ.py
from io import BytesIO
from docx import Document
from docx.shared import Pt
import requests
import time
from datetime import datetime
from decimal import Decimal, ROUND_DOWN
import streamlit as st

# 1. st.set_page_config()
st.set_page_config(
    layout="wide",
    page_title="Simulador Pessoa Jurídica",
    page_icon="imgs/v-c.png", # Verifique o caminho
    initial_sidebar_state="expanded"
)
print("INFO_LOG (Simulador_PJ.py): st.set_page_config executado.")

# 2. BLOCO DE VERIFICAÇÃO DE AUTENTICAÇÃO
auth_status = st.session_state.get("authentication_status", False)
if auth_status is not True:
    st.error("🔒 Acesso Negado! Por favor, faça login na página principal para continuar.")
    print(f"ACCESS_DENIED_LOG (Simulador_PJ.py): User not authenticated. Status: {auth_status}")
    try:
        st.page_link("Simulador_Comercial.py", label="Ir para Login", icon="🏠") # [cite: luizbicalho2024/simulador-telemetria/Simulador-Telemetria-76eb67ea5bdfd65ec4d387b010a2c3d452d296b4/Simulador_Comercial.py]
    except AttributeError:
        st.info("Retorne à página principal para efetuar o login.")
    st.stop()

current_username = st.session_state.get('username', 'N/A')
current_role = st.session_state.get('role', 'Indefinido')
current_name = st.session_state.get('name', 'N/A')
print(f"INFO_LOG (Simulador_PJ.py): User '{current_username}' authenticated. Role: '{current_role}'.")

# 3. Restante do código da sua página

API_KEY_CLOUDCONVERT = st.secrets.get("API_KEY_CLOUDCONVERT")
api_key_presente = bool(API_KEY_CLOUDCONVERT)

if not api_key_presente:
    print("WARN_LOG (Simulador_PJ.py): CLOUDCONVERT_API_KEY não configurada.")

try:
    st.image("imgs/logo.png", width=250)
except Exception as e_img:
    print(f"WARN_LOG (Simulador_PJ.py): Erro ao carregar imgs/logo.png: {e_img}")

st.markdown("<h1 style='text-align: center; color: #54A033;'>Simulador de Venda - Pessoa Jurídica</h1>", unsafe_allow_html=True)
st.markdown("---")
st.write(f"Usuário: {current_name} ({current_username})")
st.write(f"Nível de Acesso: {current_role}")
st.markdown("---")
print("INFO_LOG (Simulador_PJ.py): Cabeçalho e informações do usuário renderizados.")

planos = {
    "12 Meses": {"GPRS / Gsm": Decimal("80.88"), "Satélite": Decimal("193.80"), "Identificador de Motorista / RFID": Decimal("19.25"), "Leitor de Rede CAN / Telemetria": Decimal("75.25"), "Videomonitoramento + DMS + ADAS": Decimal("409.11")},
    "24 Meses": {"GPRS / Gsm": Decimal("53.92"), "Satélite": Decimal("129.20"), "Identificador de Motorista / RFID": Decimal("12.83"), "Leitor de Rede CAN / Telemetria": Decimal("50.17"), "Videomonitoramento + DMS + ADAS": Decimal("272.74")},
    "36 Meses": {"GPRS / Gsm": Decimal("44.93"), "Satélite": Decimal("107.67"), "Identificador de Motorista / RFID": Decimal("10.69"), "Leitor de Rede CAN / Telemetria": Decimal("41.81"), "Videomonitoramento + DMS + ADAS": Decimal("227.28")}
}
produtos_descricao = {
    "GPRS / Gsm": "Equipamento de rastreamento GSM/GPRS 2G ou 4G", "Satélite": "Equipamento de rastreamento via satélite",
    "Identificador de Motorista / RFID": "Identificação automática de motoristas via RFID",
    "Leitor de Rede CAN / Telemetria": "Leitura de dados de telemetria via rede CAN",
    "Videomonitoramento + DMS + ADAS": "Sistema de videomonitoramento com assistência ao motorista"
}

st.sidebar.header("📝 Configurações PJ")
qtd_veiculos_key = "pj_qtd_veiculos_sb_v25"
temp_contrato_key = "pj_temp_contrato_sb_v25"
qtd_veiculos_input = st.sidebar.number_input("Quantidade de Veículos 🚗", min_value=1, value=1, step=1, key=qtd_veiculos_key)
temp_contrato_selecionado_str = st.sidebar.selectbox("Tempo de Contrato ⏳", list(planos.keys()), key=temp_contrato_key)
print("INFO_LOG (Simulador_PJ.py): Widgets da sidebar renderizados.")

st.markdown("### 🛠️ Selecione os Produtos:")
col1_pj, col2_pj = st.columns(2)
produtos_selecionados_pj = {}
for i, (produto, preco_decimal) in enumerate(planos[temp_contrato_selecionado_str].items()):
    col_target = col1_pj if i % 2 == 0 else col2_pj
    produto_toggle_key = f"pj_toggle_{temp_contrato_selecionado_str.replace(' ','_')}_{produto.replace(' ', '_').replace('/', '_').replace('+', '')}_v25"
    if col_target.toggle(f"{produto} - R$ {preco_decimal:,.2f}", key=produto_toggle_key):
        produtos_selecionados_pj[produto] = preco_decimal
print(f"DEBUG_LOG (Simulador_PJ.py): Produtos selecionados para proposta: {produtos_selecionados_pj}")

soma_mensal_produtos_selecionados_calculada = sum(produtos_selecionados_pj.values()) if produtos_selecionados_pj else Decimal("0")
qtd_veiculos_decimal = Decimal(str(qtd_veiculos_input))
valor_mensal_total_frota_calculado = soma_mensal_produtos_selecionados_calculada * qtd_veiculos_decimal
meses_contrato_decimal = Decimal(temp_contrato_selecionado_str.split()[0])
valor_total_contrato_calculado = valor_mensal_total_frota_calculado * meses_contrato_decimal

st.markdown("---")
if produtos_selecionados_pj:
    st.success(f"✅ Valor Mensal por Veículo (soma dos produtos selecionados): R$ {soma_mensal_produtos_selecionados_calculada:,.2f}")
    st.info(f"💰 Valor Mensal Total para a Frota ({qtd_veiculos_input} veíc.): R$ {valor_mensal_total_frota_calculado:,.2f}")
    st.info(f"📄 Valor Total do Contrato ({temp_contrato_selecionado_str}): R$ {valor_total_contrato_calculado:,.2f}")
else:
    st.info("Selecione produtos para ver o cálculo.")
print("INFO_LOG (Simulador_PJ.py): Seção de cálculo de totais renderizada.")

if st.button("🔄 Limpar Seleção e Recalcular", key="pj_btn_limpar_recalcular_v25"):
    print("INFO_LOG (Simulador_PJ.py): Botão 'Limpar Seleção' clicado.")
    st.rerun()

# --- FUNÇÃO AUXILIAR PARA PREENCHER O DOCX ---
def replace_text_in_paragraph_runs(paragraph, placeholder, value_str):
    """
    Tenta substituir um placeholder dentro dos runs de um parágrafo.
    Esta é uma abordagem mais robusta para tentar preservar a formatação.
    Retorna True se alguma substituição foi feita no parágrafo.
    """
    substituted_in_paragraph = False
    # Concatena o texto dos runs para encontrar o placeholder
    # Esta parte é complexa porque o placeholder pode estar dividido entre runs.
    # Para uma substituição simples, vamos verificar se o placeholder está no texto do parágrafo
    # e depois tentar substituir nos runs.
    if placeholder in paragraph.text:
        # print(f"DEBUG_LOG (replace_text_in_paragraph_runs): Placeholder '{placeholder}' detectado em '{paragraph.text[:70]}...'")
        # Tenta substituir run a run. Isso é melhor para manter a formatação,
        # mas pode falhar se o placeholder estiver quebrado entre runs.
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value_str)
                # print(f"DEBUG_LOG (replace_text_in_paragraph_runs): Substituído em run: '{placeholder}' -> '{value_str}'")
                substituted_in_paragraph = True 
                # Não dá break aqui, pois o placeholder pode aparecer múltiplas vezes ou ser parcialmente substituído.

        # Fallback: Se após a tentativa de substituição run-a-run o placeholder ainda existir
        # (o que pode acontecer se ele estiver quebrado entre runs de forma complexa),
        # tenta uma substituição no texto completo do parágrafo.
        # Isso pode perder alguma formatação específica de runs.
        if placeholder in paragraph.text and not substituted_in_paragraph: # Apenas se a substituição run a run não pegou tudo
            # print(f"DEBUG_LOG (replace_text_in_paragraph_runs): Placeholder '{placeholder}' ainda no texto do parágrafo. Tentando substituição geral.")
            # paragraph.text = paragraph.text.replace(placeholder, value_str) # Esta linha pode ser problemática com formatação
            # Uma abordagem mais segura para o fallback é reconstruir os runs se necessário,
            # mas isso é muito mais complexo. Por ora, vamos focar na substituição run a run.
            # Se a substituição run a run falhou em remover completamente o placeholder,
            # é provável que o placeholder esteja mal formatado no DOCX (dividido entre runs).
            pass # Evita a substituição de paragraph.text que pode quebrar a formatação

    return substituted_in_paragraph


def preencher_proposta_docx(doc, nome_empresa_val, nome_responsavel_val, nome_consultor_val,
                            validade_proposta_dt_val,
                            produtos_selecionados_dict,
                            produtos_descricao_dict,
                            soma_total_mensal_por_veiculo_decimal,
                            qtd_veiculos_val,
                            tempo_contrato_str_val,
                            valor_mensal_total_frota_val,
                            valor_total_do_contrato_val
                           ):
    print(f"DEBUG_LOG (preencher_proposta_docx): Iniciando preenchimento para '{nome_empresa_val}'.")

    # USE PLACEHOLDERS COMO {{CHAVE}} NO SEU DOCX
    placeholders_gerais = {
        "{{NOME_EMPRESA}}": nome_empresa_val,
        "{{NOME_RESPONSAVEL}}": nome_responsavel_val,
        "{{DATA_VALIDADE}}": validade_proposta_dt_val.strftime("%d/%m/%Y"),
        "{{NOME_CONSULTOR}}": nome_consultor_val,
        "{{QTD_VEICULOS}}": str(qtd_veiculos_val),
        "{{TEMPO_CONTRATO}}": tempo_contrato_str_val,
        "{{VALOR_MENSAL_FROTA}}": f"R$ {valor_mensal_total_frota_val:,.2f}",
        "{{VALOR_TOTAL_CONTRATO}}": f"R$ {valor_total_do_contrato_val:,.2f}"
    }
    print(f"DEBUG_LOG (preencher_proposta_docx): Placeholders a serem substituídos: {placeholders_gerais}")

    # Itera sobre parágrafos e células de tabelas para substituir placeholders gerais
    for placeholder, value in placeholders_gerais.items():
        value_str = str(value) # Garante que o valor é string
        # Parágrafos do corpo principal
        for p_idx, p in enumerate(doc.paragraphs):
            if placeholder in p.text:
                print(f"DEBUG_LOG: Placeholder '{placeholder}' encontrado no parágrafo {p_idx} do corpo.")
                # Tenta substituir em cada run do parágrafo
                for run in p.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value_str)
                # Fallback: se o placeholder ainda existir (pode estar quebrado entre runs)
                if placeholder in p.text:
                    p.text = p.text.replace(placeholder, value_str) # Isso pode afetar a formatação
                    print(f"DEBUG_LOG: Substituição de fallback no texto do parágrafo {p_idx} para '{placeholder}'.")


        # Células de todas as tabelas
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for p_in_cell_idx, p_in_cell in enumerate(cell.paragraphs):
                        if placeholder in p_in_cell.text:
                            print(f"DEBUG_LOG: Placeholder '{placeholder}' encontrado na Tabela {table_idx}, Linha {row_idx}, Célula {cell_idx}, Parágrafo {p_in_cell_idx}.")
                            for run in p_in_cell.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value_str)
                            if placeholder in p_in_cell.text:
                                p_in_cell.text = p_in_cell.text.replace(placeholder, value_str)
                                print(f"DEBUG_LOG: Substituição de fallback no texto do parágrafo da célula para '{placeholder}'.")
    print("DEBUG_LOG (preencher_proposta_docx): Substituição de placeholders gerais concluída.")


    # --- PREENCHIMENTO DA TABELA DE ITENS ---
    table_to_fill = None
    expected_headers = ["Item", "Descrição", "Preço | Mês"]
    print(f"DEBUG_LOG (preencher_proposta_docx): Cabeçalhos esperados na tabela de itens: {expected_headers}")
    table_found_and_filled_correctly = False

    for table_idx, table in enumerate(doc.tables):
        print(f"DEBUG_LOG (preencher_proposta_docx): Verificando Tabela {table_idx} para itens...")
        if len(table.rows) > 0 and len(table.columns) >= len(expected_headers):
            header_cells_text_from_doc = [cell.text.strip().lower() for cell in table.rows[0].cells[:len(expected_headers)]]
            expected_headers_lower = [h.lower() for h in expected_headers]
            
            if header_cells_text_from_doc == expected_headers_lower:
                table_to_fill = table
                print(f"INFO_LOG (preencher_proposta_docx): Tabela de itens encontrada (Índice: {table_idx}).")
                
                print(f"DEBUG_LOG (preencher_proposta_docx): Limpando linhas. Linhas antes: {len(table_to_fill.rows)}")
                for i in range(len(table_to_fill.rows) - 1, 0, -1):
                    row_to_remove = table_to_fill.rows[i]
                    table_to_fill._tbl.remove(row_to_remove._tr)
                print(f"DEBUG_LOG (preencher_proposta_docx): Linhas após limpeza: {len(table_to_fill.rows)}")

                if not produtos_selecionados_dict:
                    print("WARN_LOG (preencher_proposta_docx): Nenhum produto selecionado para tabela.")
                else:
                    print(f"DEBUG_LOG (preencher_proposta_docx): Adicionando itens: {produtos_selecionados_dict}")
                    for produto_sel, preco_sel_decimal in produtos_selecionados_dict.items():
                        row_cells = table_to_fill.add_row().cells
                        row_cells[0].text = produto_sel
                        row_cells[1].text = produtos_descricao_dict.get(produto_sel, " ") 
                        row_cells[2].text = f"R$ {preco_sel_decimal:,.2f}"
                
                total_row = table_to_fill.add_row().cells
                total_row[0].text = "Total Mensal por Veículo"; total_row[0].paragraphs[0].runs[0].font.bold = True 
                total_row[1].text = "" 
                total_row[2].text = f"R$ {soma_total_mensal_por_veiculo_decimal:,.2f}"; total_row[2].paragraphs[0].runs[0].font.bold = True
                print(f"DEBUG_LOG (preencher_proposta_docx): Linha total por veículo: R$ {soma_total_mensal_por_veiculo_decimal:,.2f}")
                table_found_and_filled_correctly = True
                break 
            else:
                 print(f"DEBUG_LOG (preencher_proposta_docx): Tabela {table_idx} não correspondeu. Encontrado: {header_cells_text_from_doc}, Esperado: {expected_headers_lower}")
    
    if not table_found_and_filled_correctly:
        print("WARN_LOG (preencher_proposta_docx): Tabela de itens principal NÃO encontrada/preenchida.")
    
    return table_found_and_filled_correctly


# --- Formulário para Gerar Proposta ---
if produtos_selecionados_pj:
    st.markdown("---")
    st.subheader("📄 Gerar Proposta")

    if not api_key_presente:
        st.warning("⚠️ Geração de PDF (CloudConvert) desativada: Chave API não configurada.")
        print("WARN_LOG (Simulador_PJ.py): Geração de PDF desativada.")

    with st.form(f"formulario_proposta_pj_v25_final", clear_on_submit=False): # clear_on_submit=False para manter dados
        form_nome_empresa = st.text_input("Nome da Empresa", key="pj_form_nome_empresa_v25_final")
        form_nome_responsavel = st.text_input("Nome do Responsável", key="pj_form_nome_responsavel_v25_final")
        form_nome_consultor = st.text_input("Nome do Consultor Comercial", value=current_name, key="pj_form_nome_consultor_v25_final")
        form_validade_proposta_dt = st.date_input("Validade da Proposta", value=datetime.today(), key="pj_form_validade_proposta_v25_final")
        
        col_btn_form1, col_btn_form2 = st.columns(2)
        with col_btn_form1:
            gerar_docx_btn = st.form_submit_button("Gerar Proposta em DOCX")
        with col_btn_form2:
            gerar_pdf_cloudconvert_btn = st.form_submit_button("Gerar PDF (CloudConvert)", disabled=(not api_key_presente))

    if gerar_docx_btn:
        print(f"INFO_LOG (Simulador_PJ.py): Botão 'Gerar DOCX' clicado. Empresa: {form_nome_empresa}")
        if not all([form_nome_empresa, form_nome_responsavel, form_nome_consultor]):
            st.warning("Preencha os dados da proposta (Empresa, Responsável, Consultor).")
        elif not produtos_selecionados_pj:
             st.warning("Nenhum produto selecionado para incluir na proposta.")
        else:
            try:
                doc_template_path = "Proposta Comercial e Intenção - Verdio.docx" # [cite: luizbicalho2024/simulador-telemetria/Simulador-Telemetria-76eb67ea5bdfd65ec4d387b010a2c3d452d296b4/pages/Proposta Comercial e Intenção - Verdio.docx]
                doc = Document(doc_template_path)
                print(f"DEBUG_LOG (Simulador_PJ.py): Carregado template DOCX: {doc_template_path}")
                
                tabela_foi_preenchida = preencher_proposta_docx(
                    doc, form_nome_empresa, form_nome_responsavel, form_nome_consultor,
                    form_validade_proposta_dt, produtos_selecionados_pj,
                    produtos_descricao, soma_mensal_produtos_selecionados_calculada,
                    qtd_veiculos_input, temp_contrato_selecionado_str,
                    valor_mensal_total_frota_calculado, valor_total_contrato_calculado
                )
                
                if not tabela_foi_preenchida:
                    st.warning("Atenção: A tabela de itens não foi encontrada/preenchida no template. Verifique os cabeçalhos do DOCX e os placeholders.")
                
                buffer_docx = BytesIO()
                doc.save(buffer_docx)
                buffer_docx.seek(0)
                
                st.download_button(
                    label="📥 Baixar Proposta em DOCX",
                    data=buffer_docx,
                    file_name=f"Proposta_Verdio_{form_nome_empresa.replace(' ', '_')}_{form_validade_proposta_dt.strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="pj_download_docx_btn_v25_final"
                )
                st.success("Proposta em DOCX pronta para download!")

            except FileNotFoundError:
                st.error(f"ERRO: O arquivo template DOCX '{doc_template_path}' não foi encontrado.")
                print(f"ERROR_LOG (Simulador_PJ.py): Template DOCX '{doc_template_path}' não encontrado.")
            except Exception as e_gerar_docx:
                st.error(f"Erro inesperado ao gerar proposta DOCX: {e_gerar_docx}")
                print(f"ERROR_LOG (Simulador_PJ.py): Erro na geração DOCX: {type(e_gerar_docx).__name__} - {e_gerar_docx}")

    if gerar_pdf_cloudconvert_btn and api_key_presente:
        # (A lógica para gerar PDF permanece a mesma da resposta anterior,
        #  chamando a mesma função preencher_proposta_docx)
        # ... (código para CloudConvert) ...
        print(f"INFO_LOG (Simulador_PJ.py): Botão 'Gerar PDF (CloudConvert)' clicado. Empresa: {form_nome_empresa}")
        if not all([form_nome_empresa, form_nome_responsavel, form_nome_consultor]):
            st.warning("Preencha os dados da proposta (Empresa, Responsável, Consultor).")
        elif not produtos_selecionados_pj:
             st.warning("Nenhum produto selecionado para incluir na proposta PDF.")
        else:
            try:
                doc_template_path = "Proposta Comercial e Intenção - Verdio.docx" # [cite: luizbicalho2024/simulador-telemetria/Simulador-Telemetria-76eb67ea5bdfd65ec4d387b010a2c3d452d296b4/pages/Proposta Comercial e Intenção - Verdio.docx]
                doc = Document(doc_template_path)
                
                tabela_foi_preenchida_pdf = preencher_proposta_docx(
                    doc, form_nome_empresa, form_nome_responsavel, form_nome_consultor, 
                    form_validade_proposta_dt, produtos_selecionados_pj, 
                    produtos_descricao, soma_mensal_produtos_selecionados_calculada,
                    qtd_veiculos_input, temp_contrato_selecionado_str,
                    valor_mensal_total_frota_calculado, valor_total_contrato_calculado
                )

                if not tabela_foi_preenchida_pdf:
                    st.warning("Atenção: A tabela de itens não foi encontrada no template. PDF pode estar incompleto.")
                
                buffer_docx_for_pdf = BytesIO()
                doc.save(buffer_docx_for_pdf)
                buffer_docx_for_pdf.seek(0)
                print("INFO_LOG (Simulador_PJ.py): DOCX gerado para conversão PDF.")

                with st.spinner("Gerando PDF da proposta via CloudConvert, aguarde..."):
                    headers = {"Authorization": f"Bearer {API_KEY_CLOUDCONVERT}"} 
                    job_payload = {
                        "tasks": {
                            "import-docx": {"operation": "import/upload", "filename": f"proposta_{form_nome_empresa.replace(' ', '_')}.docx"},
                            "convert-to-pdf": {"operation": "convert", "input": "import-docx", "input_format": "docx", "output_format": "pdf", "engine": "libreoffice"},
                            "export-pdf": {"operation": "export/url", "input": "convert-to-pdf", "inline": False, "archive_multiple_files": False}
                        }
                    }
                    job_creation_response = requests.post('https://api.cloudconvert.com/v2/jobs', json=job_payload, headers=headers)
                    job_creation_response.raise_for_status() 
                    job_data = job_creation_response.json()
                    
                    upload_task_id = None; upload_url = None; upload_parameters = None
                    for task_details in job_data['data']['tasks']:
                        if task_details['operation'] == 'import/upload':
                            upload_task_id = task_details['id']
                            if task_details.get('result') and task_details['result'].get('form'):
                                upload_url = task_details['result']['form'].get('url')
                                upload_parameters = task_details['result']['form'].get('parameters')
                            break 
                    
                    if not upload_url or upload_parameters is None:
                        st.error("Falha ao obter URL/parâmetros de upload do CloudConvert.")
                        st.stop() 
                    
                    files_payload_for_upload = {'file': (f"proposta_{form_nome_empresa.replace(' ', '_')}.docx", buffer_docx_for_pdf, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
                    upload_post_response = requests.post(upload_url, data=upload_parameters, files=files_payload_for_upload)
                    upload_post_response.raise_for_status()

                    job_id = job_data['data']['id']
                    status_check_url = f'https://api.cloudconvert.com/v2/jobs/{job_id}'
                    
                    max_wait_time_seconds = 120 
                    start_process_time = time.time()
                    pdf_download_url = None
                    final_job_status = None
                    with st.empty(): 
                        while time.time() - start_process_time < max_wait_time_seconds:
                            st.write(f"Verificando status da conversão (Job ID: {job_id})...")
                            job_status_response = requests.get(status_check_url, headers=headers)
                            job_status_response.raise_for_status()
                            job_status_data = job_status_response.json()
                            current_job_status = job_status_data['data']['status']
                            print(f"DEBUG_LOG (Simulador_PJ.py): Status job CloudConvert '{job_id}': {current_job_status}")
                            
                            if current_job_status == 'finished':
                                final_job_status = 'finished'
                                for task_data_loop in job_status_data['data']['tasks']:
                                    if task_data_loop.get('operation') == 'export/url' and task_data_loop.get('status') == 'finished': 
                                        if task_data_loop.get('result') and task_data_loop['result'].get('files') and len(task_data_loop['result']['files']) > 0:
                                            pdf_download_url = task_data_loop['result']['files'][0]['url']
                                            break 
                                if pdf_download_url: break 
                            elif current_job_status == 'error':
                                final_job_status = 'error'
                                # (lógica de extração de erro)
                                st.error(f"Erro na conversão PDF via CloudConvert.")
                                break 
                            time.sleep(4) 
                        
                        if pdf_download_url:
                            pdf_file_content = requests.get(pdf_download_url).content 
                            st.download_button(
                                label="📥 Baixar Proposta em PDF",
                                data=pdf_file_content,
                                file_name=f"Proposta_Verdio_{form_nome_empresa.replace(' ', '_')}_{form_validade_proposta_dt.strftime('%Y%m%d')}.pdf",
                                mime="application/pdf",
                                key="pj_download_pdf_btn_v25_final_cc" 
                            )
                        elif final_job_status != 'error': 
                            st.error("Não foi possível obter o PDF ou tempo de espera excedido.")

            except FileNotFoundError:
                st.error(f"ERRO: Template DOCX '{doc_template_path}' não encontrado.")
            except requests.exceptions.RequestException as req_e:
                st.error(f"Erro de comunicação com CloudConvert: {req_e}")
            except Exception as e_gerar_pdf:
                st.error(f"Erro inesperado ao gerar proposta PDF: {e_gerar_pdf}")
                print(f"ERROR_LOG (Simulador_PJ.py): Erro na geração PDF: {type(e_gerar_pdf).__name__} - {e_gerar_pdf}")
    
elif not api_key_presente and produtos_selecionados_pj: 
    pass 
elif not produtos_selecionados_pj: 
    st.info("Selecione produtos para preencher dados e gerar uma proposta.")

print("INFO_LOG (Simulador_PJ.py): Renderização da página PJ concluída.")
